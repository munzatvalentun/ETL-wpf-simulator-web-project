from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── сторінка ────────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(1.5)
sec.top_margin    = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.page_width    = Cm(21.0)
sec.page_height   = Cm(29.7)

TNR  = 'Times New Roman'
CONS = 'Consolas'

# ── хелпери ──────────────────────────────────────────────────────────────────
def _fix_font(run, name):
    rPr = run._r.get_or_add_rPr()
    rf  = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rPr.insert(0, rf)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rf.set(qn(attr), name)

def run_font(run, name=TNR, size=14, bold=False, italic=False, rgb=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if rgb:
        run.font.color.rgb = RGBColor(*rgb)
    _fix_font(run, name)

def _para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              first_cm=1.25, sb=0, sa=6, left_cm=0):
    pf = p.paragraph_format
    pf.alignment          = align
    pf.line_spacing_rule  = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent  = Cm(first_cm)
    pf.space_before       = Pt(sb)
    pf.space_after        = Pt(sa)
    if left_cm:
        pf.left_indent = Cm(left_cm)

def body(text, bold=False, italic=False, indent=True):
    p = doc.add_paragraph()
    _para_fmt(p, first_cm=1.25 if indent else 0)
    r = p.add_run(text)
    run_font(r, bold=bold, italic=italic)
    return p

def h1(text):
    p = doc.add_paragraph()
    _para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_cm=0, sb=18, sa=12)
    r = p.add_run(text.upper())
    run_font(r, bold=True)
    return p

def h2(text):
    p = doc.add_paragraph()
    _para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_cm=0, sb=12, sa=6)
    r = p.add_run(text)
    run_font(r, bold=True)
    return p

def h3(text):
    p = doc.add_paragraph()
    _para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_cm=0, sb=8, sa=4)
    r = p.add_run(text)
    run_font(r, bold=True, italic=True)
    return p

def fig_placeholder(num, desc):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 0)
    trPr = t.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement('w:trHeight')
    th.set(qn('w:val'), '1400')
    th.set(qn('w:hRule'), 'atLeast')
    trPr.append(th)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(f'[ ВСТАВИТИ РИСУНОК {num}: {desc} ]')
    run_font(r, size=11, italic=True, rgb=(0x88, 0x88, 0x88))
    cap = doc.add_paragraph()
    _para_fmt(cap, align=WD_ALIGN_PARAGRAPH.CENTER, first_cm=0, sb=2, sa=10)
    rc = cap.add_run(f'Рисунок {num} — {desc}')
    run_font(rc, size=12, italic=True)

def code_line(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Cm(0)
    pf.left_indent        = Cm(0.5)
    pf.space_before       = Pt(0)
    pf.space_after        = Pt(0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = CONS
    r.font.size = Pt(9)
    _fix_font(r, CONS)

def code_block(src):
    for line in src.split('\n'):
        code_line(line)
    gap = doc.add_paragraph()
    gap.paragraph_format.space_after  = Pt(8)
    gap.paragraph_format.space_before = Pt(0)

def note(text):
    p = doc.add_paragraph()
    _para_fmt(p, first_cm=0, left_cm=1.0, sb=0, sa=4)
    r = p.add_run(text)
    run_font(r, size=12, italic=True, rgb=(0x44, 0x44, 0x44))

# ════════════════════════════════════════════════════════════════════════════
#  РОЗДІЛ 3
# ════════════════════════════════════════════════════════════════════════════

h1('РОЗДІЛ 3\nОСОБЛИВОСТІ ПРОГРАМНОЇ РЕАЛІЗАЦІЇ СИСТЕМИ')

# ── 3.1 ─────────────────────────────────────────────────────────────────────
h2('3.1. Загальна архітектура рішення')

body(
    'Розроблена система складається з двох самостійних застосунків, '
    'що взаємодіють через спільну реляційну базу даних: '
    'десктопного WPF-симулятора (ETL_simulator) та веб-застосунку '
    'моніторингу (ETL_web_project). Такий поділ відповідає принципу '
    'розмежування відповідальностей: симулятор відповідає за генерацію '
    'даних і виконання ETL-пайплайну, а веб-застосунок — '
    'за спостереження, аналітику та адміністрування.'
)
body(
    'Проект ETL_simulator посилається на ETL_web_project, '
    'тому спільні сутності (Entity-класи, DbContext, перерахування) '
    'визначено лише в одному місці та повторно використовуються обома '
    'застосунками без дублювання. '
    'База даних EtlDb розгорнута на Microsoft SQL Server; '
    'підключення налаштовується у файлі appsettings.json кожного проекту.'
)
body(
    'Технологічний стек: .NET 8, WPF (MVVM), ASP.NET Core 8 MVC, '
    'Entity Framework Core 8 (Code First), бібліотека Bogus для генерації '
    'тестових даних, AutoMapper для перетворення сутностей на DTO.'
)

fig_placeholder('3.1', 'Схема архітектури рішення: два застосунки, спільна БД, Named Pipe')

# ── 3.2 ─────────────────────────────────────────────────────────────────────
h2('3.2. Структура бази даних')

body(
    'База даних організована за принципом схемного розподілу, '
    'що безпосередньо відображає шари обробки даних. '
    'Кожна SQL-схема ізолює таблиці певного рівня, '
    'що унеможливлює випадкове змішування даних різного ступеня якості '
    'та спрощує керування правами доступу.'
)

body('Схемна структура включає такі рівні:', indent=False)
for row in [
    ('stg',    'SalesRaw',                     'сирі вхідні записи, що ще не пройшли валідацію'),
    ('silver', 'SalesClean',                   'очищені та дедубліковані записи'),
    ('dw',     'FactSales, DimStore, DimProduct, DimDate, DimCustomer', 'сховище даних зіркоподібної схеми'),
    ('etl',    'EtlJob, EtlRun, EtlLog, EtlSchedule', 'метадані управління пайплайном'),
    ('auth',   'UserAccount',                  'облікові записи з унікальністю за username та email'),
]:
    p = doc.add_paragraph(style='List Bullet')
    _para_fmt(p, first_cm=0, left_cm=1.25, sb=0, sa=2)
    r1 = p.add_run(f'{row[0]}.{row[1]}')
    run_font(r1, bold=True, size=13)
    r2 = p.add_run(f' — {row[2]}')
    run_font(r2, size=13)

body(
    'Схема dw реалізує зіркоподібну (star schema) модель: '
    'центральна таблиця фактів FactSales зберігає кількісні показники '
    '(Quantity, TotalAmount) та зовнішні ключі до чотирьох вимірів. '
    'Виміри наповнюються автоматично під час Gold-фази ETL. '
    'Таблиця DimDate заповнюється динамічно: '
    'якщо запис із датою ще не існує, він створюється '
    'зі всіма похідними атрибутами (рік, місяць, назва дня тижня тощо).'
)
body(
    'Структура бази даних підтримується засобами Entity Framework Core '
    'у режимі Code First: усі таблиці, зв\'язки та обмеження описані '
    'через C#-класи та FluentAPI у ProjectContext.OnModelCreating. '
    'Еволюція схеми відстежується міграціями EF Core; '
    'на момент завершення розробки створено дванадцять міграцій.'
)

fig_placeholder('3.2', 'ER-діаграма бази даних EtlDb (усі схеми)')

# ── 3.3 ─────────────────────────────────────────────────────────────────────
h2('3.3. Реалізація ETL-пайплайну за Medallion Architecture')

body(
    'ETL-пайплайн реалізує патерн Medallion Architecture — '
    'триступеневу модель обробки даних Bronze → Silver → Gold. '
    'Кожен ступінь відповідає окремому C#-класу з чіткою '
    'відповідальністю, що забезпечує незалежність тестування '
    'та можливість заміни реалізації без впливу на решту пайплайну.'
)

h3('3.3.1. Bronze-шар: завантаження сирих даних')
body(
    'Клас BronzeLoader реалізує найпростіший можливий контракт: '
    'прийняти колекцію об\'єктів SalesRaw та зберегти їх у таблиці stg.SalesRaw. '
    'На цьому шарі дані не перевіряються — завдання Bronze полягає '
    'виключно у збереженні всього, що надійшло від генератора, '
    'включно з потенційно дефектними записами. '
    'Це дає змогу у будь-який момент переаналізувати вхідні дані '
    'без їх повторної генерації.'
)

h3('3.3.2. Silver-шар: очищення та валідація')
body(
    'Клас SilverProcessor є центральним елементом якісного контролю. '
    'Він вибирає з Bronze усі необроблені записи '
    '(прапорець IsProcessedToSilver = false) і послідовно застосовує '
    'правила валідації: перевірка наявності дати продажу, '
    'коду магазину, коду продукту, а також позитивності '
    'кількості та ціни. '
    'Додатково виконується двоетапна дедублікація: '
    'спочатку в межах поточного батчу через HashSet, '
    'а потім — запитом до існуючих записів silver.SalesClean.'
)
body(
    'Ключовою особливістю реалізації є те, що кожен відхилений запис '
    'отримує конкретну причину відмови, яка зберігається у словнику '
    'RejectedReasons (Dictionary<int, string>, де ключ — ідентифікатор '
    'запису SalesRaw). Це дає можливість WPF-симулятору відображати '
    'причину відхилення в анімованій візуалізації, '
    'а EtlOrchestrator — формувати деталізований лог у базі даних. '
    'Метод повертає запис SilverResult, '
    'що містить лічильники по кожному типу відхилення '
    'та словник причин для анімації.'
)

h3('3.3.3. Gold-шар: завантаження до сховища даних')
body(
    'Клас GoldLoader реалізує збагачення даних при переміщенні '
    'з silver.SalesClean до dw.FactSales. '
    'Для кожного очищеного запису виконується пошук відповідних '
    'записів у вимірах DimStore та DimProduct. '
    'Якщо відповідного виміру не знайдено, '
    'запис пропускається без помилки — '
    'це захист від ситуацій, коли виміри ще не заповнені. '
    'Окремий метод GetOrCreateDateKeyAsync забезпечує ідемпотентне '
    'створення записів у DimDate: якщо запис з потрібною датою '
    'вже існує, повертається наявний ключ; '
    'у протилежному випадку запис створюється з усіма '
    'похідними атрибутами часу.'
)
body(
    'Метод SeedDimensionsAsync викликається одноразово при запуску '
    'симулятора: він заповнює виміри DimStore, DimProduct та DimCustomer '
    'з пулів, визначених у класі SalesGenerator, '
    'пропускаючи вже існуючі записи. '
    'Це гарантує консистентність вимірів '
    'незалежно від кількості запусків симулятора.'
)

h3('3.3.4. Оркестрація: EtlOrchestrator')
body(
    'Клас EtlOrchestrator відповідає за управління метаданими '
    'виконання пайплайну: він не містить бізнес-логіки ETL, '
    'але забезпечує повний цикл аудиту. '
    'Метод StartRunAsync знаходить або створює запис EtlJob, '
    'потім створює новий EtlRun зі статусом Running. '
    'Метод FinishRunAsync закриває запис з відповідним статусом '
    '(Success або Failed) та часом завершення. '
    'Метод LogAsync додає запис до etl.EtlLog з рівнями '
    'Info, Warn або Error.'
)
body(
    'Принципово важливим є те, що оркестратор не викликається '
    'всередині BronzeLoader, SilverProcessor або GoldLoader. '
    'Усі звернення до EtlOrchestrator зосереджені у MainViewModel, '
    'що є єдиним координатором кроку пайплайну. '
    'Це дозволяє записувати в лог стан кожного шару окремо: '
    'Bronze завжди генерує Info-лог, '
    'Silver — Info або Warn (з переліком причин відхилень), '
    'Gold — Info або Error.'
)

fig_placeholder('3.3', 'Послідовність виконання одного кроку ETL-пайплайну')

# ── 3.4 ─────────────────────────────────────────────────────────────────────
h2('3.4. WPF-симулятор: архітектура та реалізація')

h3('3.4.1. Патерн MVVM та архітектура UI')
body(
    'WPF-застосунок побудовано на основі патерну MVVM '
    '(Model-View-ViewModel). '
    'Єдиний ViewModel — клас MainViewModel — '
    'реалізує інтерфейс INotifyPropertyChanged для двостороннього '
    'зв\'язування даних з UI. '
    'Команди (Start, Stop, Step, ClearLog) реалізовані через '
    'клас RelayCommand, що інкапсулює делегати Execute та CanExecute. '
    'Завдяки цьому ViewModel не має жодної прямої залежності '
    'від класів WPF — лише від стандартних інтерфейсів '
    'System.Windows.Input та System.ComponentModel.'
)
body(
    'Для відображення журналу подій використовується '
    'ObservableCollection<LogEntry> — колекція, '
    'зміни якої автоматично відображаються в ItemsControl '
    'без додаткового коду в Code-Behind. '
    'Кольорова схема записів журналу (сірий / жовтий / червоний) '
    'задається через властивість Color типу LogEntry '
    'та прив\'язується до DataTemplate. '
    'Аналогічно реалізовано кольорове кодування поточного запису '
    'в анімаційному блоці: XAML-тригер DataTrigger '
    'перемикає фон та колір тексту залежно від значення '
    'рядкової властивості RecStatus ("passed" / "rejected").'
)

h3('3.4.2. Генератор синтетичних даних')
body(
    'Клас SalesGenerator використовує бібліотеку Bogus — '
    'інструмент для генерації реалістичних фіктивних даних. '
    'При ініціалізації статичних полів (у статичному конструкторі) '
    'формуються пули: 10 магазинів, 100 продуктів та 250 клієнтів. '
    'Статична природа пулів гарантує їхню незмінність протягом сесії '
    'та консистентність з вимірами в БД, '
    'що заповнюються методом SeedDimensionsAsync.'
)
body(
    'Клас GeneratorConfig зберігає параметри, '
    'що регулюють рівень "забрудненості" генерованих даних: '
    'NullStoreRate, NullProductRate (частка null-значень), '
    'NegativeQuantityRate (частка від\'ємних кількостей), '
    'DuplicateRate (ймовірність дублювання попереднього запису). '
    'Ці параметри прив\'язані до слайдерів у WPF-інтерфейсі '
    'та оновлюються в реальному часі. '
    'Такий підхід дозволяє моделювати різні сценарії якості '
    'вхідних даних без перезапуску симулятора.'
)

h3('3.4.3. Асинхронна логіка виконання та анімація')
body(
    'Цикл симуляції реалізований через методи RunLoopAsync '
    'та RunStepAsync, що приймають CancellationToken. '
    'Цикл запускається у фоновому потоці через Task.Run; '
    'зупинка відбувається через скасування токена без блокування UI. '
    'Кожен крок послідовно виконує Bronze, Silver та Gold, '
    'після чого передає управління методу AnimateBatchAsync.'
)
body(
    'Анімація реалізована без використання WPF Storyboard. '
    'Після завершення ETL-обробки батчу AnimateBatchAsync '
    'послідовно проходить по кожному запису '
    'з затримкою DelayMs / BatchSize мілісекунд, '
    'де DelayMs задається слайдером. '
    'Для кожного запису перевіряється його наявність у словнику '
    'RejectedReasons; якщо запис відхилений — '
    'встановлюється RecStatus = "rejected" і заповнюється RecReason. '
    'Оновлення властивостей ViewModel відбувається через '
    'Dispatcher.Invoke для безпечного звернення до UI-потоку. '
    'Завдяки цьому підходу UserControl не містить жодної логіки — '
    'лише декларативні DataTrigger у XAML.'
)

fig_placeholder('3.4', 'Вікно WPF-симулятора: pipeline visualizer та анімований запис')

# ── 3.5 ─────────────────────────────────────────────────────────────────────
h2('3.5. Веб-застосунок: архітектура та ключові модулі')

h3('3.5.1. Шаблон MVC та шар сервісів')
body(
    'Веб-застосунок побудовано на ASP.NET Core 8 MVC. '
    'Контролери (EtlController, DashboardController, '
    'AccountController, AdminController, EtlScheduleController) '
    'є тонким шаром: вони не містять бізнес-логіки, '
    'а лише делегують виклики до відповідних сервісів '
    'та передають результат до View у вигляді DTO-об\'єктів.'
)
body(
    'Кожен сервіс визначений через інтерфейс '
    '(IEtlLogService, IEtlJobService, IDashboardService тощо) '
    'і реєструється у контейнері залежностей у Program.cs. '
    'Такий підхід дозволяє замінювати реалізації сервісів '
    'без зміни коду контролерів. '
    'Для перетворення сутностей на DTO використовується AutoMapper: '
    'профілі маппінгу зосереджені у папці Mappings, '
    'що забезпечує централізований контроль над трансформацією даних.'
)

h3('3.5.2. Моніторинг ETL: Logs та Jobs')
body(
    'Сервіс EtlLogService надає аналітичний огляд журналу виконання. '
    'Метод GetLogsAsync приймає фільтри '
    '(діапазон дат, рівень, текст пошуку) '
    'та повертає об\'єкт EtlLogSummaryDto, '
    'що містить не лише відфільтровані записи, '
    'а й агреговані лічильники: '
    'загальну кількість Info/Warn/Error '
    'та окремо — кількість за останні 24 години. '
    'Ці лічильники відображаються у вигляді KPI-карток '
    'на сторінці /Etl/Logs.'
)
body(
    'Сторінка /Etl/Jobs відображає список ETL-задач '
    'з інформацією про останній запуск кожної '
    '(статус, час початку/завершення, кількість оброблених рядків). '
    'Кнопка "Run Now" надсилає HTTP POST-запит до EtlController.RunJob, '
    'який через EtlJobService намагається підключитися '
    'до Named Pipe WPF-симулятора. '
    'Результат операції (успіх або повідомлення про помилку) '
    'передається через TempData та відображається '
    'у вигляді сповіщення на сторінці.'
)

h3('3.5.3. Авторизація та ролева модель')
body(
    'Аутентифікація реалізована на основі cookie-сесій ASP.NET Core. '
    'Таблиця auth.UserAccount зберігає хешований пароль '
    'та роль користувача (Admin, DataEngineer, Viewer). '
    'Доступ до чутливих операцій обмежується атрибутом '
    '[Authorize(Roles = "Admin,DataEngineer")] '
    'на рівні методів контролерів. '
    'Сторінки адміністрування (AdminController) '
    'доступні лише для ролі Admin.'
)

fig_placeholder('3.5', 'Веб-дашборд: сторінка ETL Logs з KPI-картками')

# ── 3.6 ─────────────────────────────────────────────────────────────────────
h2('3.6. Міжпроцесна взаємодія через Named Pipe')

body(
    'Веб-застосунок і WPF-симулятор є окремими процесами '
    'без спільного адресного простору. '
    'Виклик логіки симулятора безпосередньо з веб-коду неможливий '
    'через однонаправленість залежності проектів: '
    'ETL_simulator посилається на ETL_web_project, '
    'але не навпаки (що уникає циклічних залежностей). '
    'Для реалізації кнопки "Run Now" обрано механізм '
    'Windows Named Pipe як найлегший і найнативніший '
    'варіант IPC для двох локальних Windows-процесів.'
)
body(
    'При запуску WPF-застосунку клас PipeListener '
    'стартує фоновий Task, що очікує підключення '
    'до іменованого каналу "etl-wpf-simulator". '
    'Очікування реалізовано через WaitForConnectionAsync '
    'з передачею CancellationToken, '
    'що отримує скасування при закритті застосунку (OnExit). '
    'Після отримання команди "RUN" викликається '
    'Dispatcher.Invoke для безпечного переходу на UI-поток, '
    'де викликається MainViewModel.TriggerExternalRun.'
)
body(
    'На стороні веб-застосунку EtlJobService '
    'відкриває NamedPipeClientStream з таймаутом 2 секунди. '
    'Якщо WPF-симулятор не запущено, '
    'ConnectAsync генерує TimeoutException, '
    'яке перехоплюється та перетворюється на '
    'зрозуміле повідомлення для користувача. '
    'Якщо з\'єднання успішне — надсилається рядок "RUN", '
    'після чого WPF виконує один повний крок пайплайну '
    'з поточними налаштуваннями.'
)

fig_placeholder('3.6', 'Схема взаємодії через Named Pipe: веб → WPF')

body(
    'Перевагою цього підходу є збереження існуючої архітектури '
    'без введення нових таблиць, черг або зовнішніх брокерів. '
    'Обидва застосунки залишаються незалежними: '
    'веб-застосунок не знає деталей ETL-логіки, '
    'а WPF не залежить від веб-стеку.'
)

# ════════════════════════════════════════════════════════════════════════════
#  ВИСНОВКИ ДО РОЗДІЛУ 3
# ════════════════════════════════════════════════════════════════════════════

h2('Висновки до розділу 3')

body(
    'У третьому розділі розглянуто ключові аспекти програмної '
    'реалізації системи. Архітектура двох незалежних застосунків '
    'зі спільною базою даних забезпечує чітке розмежування '
    'між виконанням ETL і його моніторингом. '
    'Схемний поділ бази даних безпосередньо відображає '
    'шари Medallion Architecture, '
    'що робить структуру даних самодокументованою. '
    'ETL-пайплайн реалізовано через три незалежних класи-обробники '
    'та клас-оркестратор, що управляє аудитом виконання. '
    'WPF-застосунок дотримується патерну MVVM та '
    'використовує таймер-анімацію для наочної демонстрації '
    'потоку даних крізь шари. '
    'Веб-застосунок побудовано на сервісному шарі '
    'з інтерфейсами та DI-контейнером. '
    'Зв\'язок між процесами реалізовано через Windows Named Pipe '
    'без зміни архітектурних залежностей між проектами.'
)

# ════════════════════════════════════════════════════════════════════════════
#  ДОДАТКИ
# ════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
h1('ДОДАТКИ')

# ── Додаток А ────────────────────────────────────────────────────────────────
h2('Додаток А. SilverProcessor — валідація та дедублікація')

note('Файл: ETL_simulator/ETL/SilverProcessor.cs (метод ProcessAsync, скорочено)')

code_block(
"""public async Task<SilverResult> ProcessAsync()
{
    var raw = await _db.SalesRaws
        .Where(r => !r.IsProcessedToSilver).ToListAsync();

    if (raw.Count == 0)
        return new SilverResult(0,0,0,0,0,0,0,0, new Dictionary<int,string>());

    var toInsert        = new List<SalesClean>();
    var seen            = new HashSet<(DateTime, string, string)>();
    var rejectedReasons = new Dictionary<int, string>();
    int rejected=0, duplicates=0,
        nullStore=0, nullProduct=0, badQty=0, badPrice=0, nullDate=0;

    foreach (var r in raw)
    {
        r.IsProcessedToSilver = true;

        // ── валідація ──────────────────────────────────────────────────────
        string? reason = null;
        if      (!r.SalesTime.HasValue)                        { reason="Відсутня дата";    nullDate++;    }
        else if (r.StoreCode   == null)                        { reason="Null StoreCode";   nullStore++;   }
        else if (r.ProductCode == null)                        { reason="Null ProductCode"; nullProduct++; }
        else if (!r.Quantity.HasValue  || r.Quantity  <= 0)   { reason="Кількість ≤ 0";   badQty++;      }
        else if (!r.UnitPrice.HasValue || r.UnitPrice <= 0)   { reason="Ціна ≤ 0";        badPrice++;    }

        if (reason != null)
            { rejected++; rejectedReasons[r.Id] = reason; continue; }

        // ── дедублікація (батч) ────────────────────────────────────────────
        var key = (r.SalesTime!.Value, r.StoreCode!, r.ProductCode!);
        if (seen.Contains(key))
            { duplicates++; rejected++; rejectedReasons[r.Id]="Дублікат (батч)"; continue; }

        // ── дедублікація (БД) ──────────────────────────────────────────────
        var exists = await _db.SalesCleans.AnyAsync(s =>
            s.SalesTime==r.SalesTime.Value &&
            s.StoreCode==r.StoreCode &&
            s.ProductCode==r.ProductCode);

        if (exists)
            { duplicates++; rejected++; rejectedReasons[r.Id]="Дублікат (БД)"; continue; }

        seen.Add(key);
        toInsert.Add(new SalesClean { /* … маппінг полів … */ });
    }

    if (toInsert.Count > 0) _db.SalesCleans.AddRange(toInsert);
    await _db.SaveChangesAsync();

    return new SilverResult(toInsert.Count, rejected, duplicates,
        nullStore, nullProduct, badQty, badPrice, nullDate, rejectedReasons);
}""")

# ── Додаток Б ────────────────────────────────────────────────────────────────
h2('Додаток Б. EtlOrchestrator — управління циклом виконання')

note('Файл: ETL_simulator/ETL/EtlOrchestrator.cs')

code_block(
"""public class EtlOrchestrator
{
    private readonly ProjectContext _db;

    public async Task<EtlRun> StartRunAsync(int rowsRead)
    {
        var job = await EnsureJobAsync();
        var run = new EtlRun
        {
            JobId     = job.JobId,
            Status    = EtlStatus.Running,
            StartTime = DateTime.Now,
            RowsRead  = rowsRead
        };
        _db.EtlRuns.Add(run);
        await _db.SaveChangesAsync();
        return run;
    }

    public async Task FinishRunAsync(EtlRun run, int rowsInserted,
                                     EtlStatus status, string error = "")
    {
        run.Status       = status;
        run.EndTime      = DateTime.Now;
        run.RowsInserted = rowsInserted;
        run.ErrorMessage = error;
        await _db.SaveChangesAsync();
    }

    public async Task LogAsync(long runId, LogLevel level, string message)
    {
        _db.EtlLogs.Add(new EtlLog
            { RunId=runId, Level=level, Message=message, LogTime=DateTime.Now });
        await _db.SaveChangesAsync();
    }

    private async Task<EtlJob> EnsureJobAsync()
    {
        var job = await _db.EtlJobs
            .FirstOrDefaultAsync(j => j.JobCode == "SIMULATOR");
        if (job != null) return job;
        job = new EtlJob
        {
            JobName="ETL Simulator", JobCode="SIMULATOR", IsActive=true
        };
        _db.EtlJobs.Add(job);
        await _db.SaveChangesAsync();
        return job;
    }
}""")

# ── Додаток В ────────────────────────────────────────────────────────────────
h2('Додаток В. MainViewModel — асинхронний крок пайплайну та анімація')

note('Файл: ETL_simulator/ViewModels/MainViewModel.cs (методи RunStepAsync та AnimateBatchAsync, скорочено)')

code_block(
"""private async Task RunStepAsync(CancellationToken ct)
{
    // ── Bronze ────────────────────────────────────────────────────────────
    List<SalesRaw> batch;
    BronzeResult br;
    try
    {
        batch = generator.GenerateBatch(BatchSize);
        br    = await bronze.LoadAsync(batch);
        // … оновлення лічильників UI, AddLog …
    }
    catch (Exception ex)
    {
        AddLog(LogLevel.Error, $"[Bronze] {GetDeepMessage(ex)}"); return;
    }

    var run = await orchestrator.StartRunAsync(br.Inserted);
    await orchestrator.LogAsync(run.RunId, DbLogLevel.Info,
        $"Bronze: завантажено {br.Inserted} сирих записів");

    // ── Silver ────────────────────────────────────────────────────────────
    SilverResult? sr = null;
    try
    {
        sr = await silver.ProcessAsync();
        // … оновлення лічильників, деталізований Warning-лог …
        await orchestrator.LogAsync(run.RunId,
            sr.Rejected > 0 ? DbLogLevel.Warn : DbLogLevel.Info,
            $"Silver: прийнято {sr.Inserted}, відхилено {sr.Rejected}");
    }
    catch (Exception ex)
    {
        await orchestrator.LogAsync(run.RunId, DbLogLevel.Error, ex.Message);
        await orchestrator.FinishRunAsync(run, 0, EtlStatus.Failed, ex.Message);
        return;
    }

    // ── Gold ──────────────────────────────────────────────────────────────
    try
    {
        var gr = await gold.LoadAsync();
        await orchestrator.LogAsync(run.RunId, DbLogLevel.Info,
            $"Gold: завантажено {gr.Inserted} до FactSales");
        await orchestrator.FinishRunAsync(run, gr.Inserted, EtlStatus.Success);
    }
    catch (Exception ex)
    {
        await orchestrator.FinishRunAsync(run, 0, EtlStatus.Failed, ex.Message);
    }

    await AnimateBatchAsync(batch, sr.RejectedReasons, ct);
}

// ─────────────────────────────────────────────────────────────────────────────

private async Task AnimateBatchAsync(List<SalesRaw> batch,
    IReadOnlyDictionary<int, string> rejectedReasons, CancellationToken ct)
{
    int msPerRecord = Math.Max(50, DelayMs / batch.Count);

    for (int i = 0; i < batch.Count; i++)
    {
        if (ct.IsCancellationRequested) break;
        var r        = batch[i];
        bool rejected = rejectedReasons.ContainsKey(r.Id);

        App.Current.Dispatcher.Invoke(() =>
        {
            RecStore    = r.StoreCode    ?? "—";
            RecProduct  = r.ProductCode  ?? "—";
            RecQty      = r.Quantity?.ToString()      ?? "—";
            RecPrice    = r.UnitPrice?.ToString("F2") ?? "—";
            RecStatus   = rejected ? "rejected" : "passed";
            RecReason   = rejected ? rejectedReasons[r.Id] : "";
            RecProgress = $"{i+1} / {batch.Count}";
            ShowCurrentRecord = true;
        });

        try { await Task.Delay(msPerRecord, ct); }
        catch (OperationCanceledException) { break; }
    }
    App.Current.Dispatcher.Invoke(() => ShowCurrentRecord = false);
}""")

# ── Додаток Г ────────────────────────────────────────────────────────────────
h2('Додаток Г. PipeListener та TriggerExternalRun — Named Pipe IPC')

note('Файли: ETL_simulator/PipeListener.cs та фрагмент MainViewModel.cs')

code_block(
"""// ── PipeListener.cs ───────────────────────────────────────────────────────
public class PipeListener
{
    public const string PipeName = "etl-wpf-simulator";
    private readonly Action _onTrigger;

    public void Start(CancellationToken ct) =>
        Task.Run(() => ListenAsync(ct), ct);

    private async Task ListenAsync(CancellationToken ct)
    {
        while (!ct.IsCancellationRequested)
        {
            try
            {
                using var server = new NamedPipeServerStream(
                    PipeName, PipeDirection.In, 1,
                    PipeTransmissionMode.Byte, PipeOptions.Asynchronous);

                await server.WaitForConnectionAsync(ct);

                using var reader = new StreamReader(server);
                if (await reader.ReadLineAsync() == "RUN")
                    App.Current.Dispatcher.Invoke(_onTrigger);
            }
            catch (OperationCanceledException) { break; }
            catch { /* ігноруємо збої каналу, продовжуємо слухати */ }
        }
    }
}

// ── MainViewModel.cs ──────────────────────────────────────────────────────
public void TriggerExternalRun()
{
    if (!IsRunning)
    {
        AddLog(LogLevel.Info, "⚡ Тригер з веб-додатку: запускаю крок...");
        _ = Task.Run(() => RunStepAsync(CancellationToken.None));
    }
}

// ── EtlJobService.cs (веб) ────────────────────────────────────────────────
public async Task<long> TriggerRunAsync(int jobId)
{
    var job = await _context.EtlJobs.FindAsync(jobId);
    if (job == null) throw new Exception("Job not found.");
    try
    {
        using var client = new NamedPipeClientStream(
            ".", "etl-wpf-simulator", PipeDirection.Out);
        await client.ConnectAsync(2000);   // таймаут 2 секунди
        using var writer = new StreamWriter(client) { AutoFlush = true };
        await writer.WriteLineAsync("RUN");
        return 0;
    }
    catch (TimeoutException)
    {
        throw new InvalidOperationException(
            "WPF Simulator не запущено. Запустіть ETL_simulator.exe.");
    }
}""")

# ════════════════════════════════════════════════════════════════════════════
out = r'C:\Git\ETL-wpf-simulator-web-project\diploma\Chapter3.docx'
doc.save(out)
print(f'Збережено: {out}')
